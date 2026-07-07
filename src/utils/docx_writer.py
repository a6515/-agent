"""
============================================================
公文 .docx 文件生成器
============================================================
将 Agent 生成的公文正文渲染为符合《党政机关公文格式》
（GB/T 9704-2012）的 .docx 文件。

格式规范：
  - 标题：   二号小标宋体（22pt），居中
  - 正文：   三号仿宋体（16pt），首行缩进 2 字符
  - 一级标题：三号黑体
  - 二级标题：三号楷体
  - 落款：   右对齐
  - 页边距： 上 3.7cm 下 3.5cm 左 2.8cm 右 2.6cm（A4）
  - 行距：   固定 28 磅
"""

import re
import os
from pathlib import Path
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from src.utils.logger import get_logger
from config.settings import settings

# 输出目录 — 统一用 settings，保证与 server.py 下载路径一致
_DATA_OUTPUT_DIR = settings.DATA_DIR / "output"

logger = get_logger(__name__)


# ============================================================
# 字体和样式常量（符合 GB/T 9704-2012）
# ============================================================

FONT_TITLE = "方正小标宋简体"   # 公文标题字体（二号）
FONT_BODY = "仿宋"              # 正文（三号）
FONT_HEITI = "黑体"             # 一级标题
FONT_KAITI = "楷体"             # 二级标题

SIZE_TITLE = Pt(22)             # 二号
SIZE_BODY = Pt(16)              # 三号

# 页边距
MARGIN_TOP = Cm(3.7)
MARGIN_BOTTOM = Cm(3.5)
MARGIN_LEFT = Cm(2.8)
MARGIN_RIGHT = Cm(2.6)


# ============================================================
# 段落解析
# ============================================================

def _classify_line(line: str, is_first_content: bool) -> str:
    """
    判断文本行在公文中的角色。

    返回:
      "title"       — 公文标题
      "subtitle"    — 发文字号/副标题
      "recipient"   — 主送机关
      "heading1"    — 一级标题（一、二、三、...）
      "heading2"    — 二级标题（（一）（二）...）
      "body"        — 正文
      "signature"   — 落款（发文机关 + 日期）
      "attachment"  — 附件说明
      "blank"       — 空行
    """
    if not line.strip():
        return "blank"

    # 附件标记
    if re.match(r"^附件[：:]", line):
        return "attachment"

    # 发文字号（如「国发〔2025〕1号」「X办发〔2024〕15号」）
    if re.match(r"^[\w一-鿿]+[〔\[（(]\s*\d{4}\s*[〕\]）)]\s*\d+\s*号", line):
        return "subtitle"

    # 落款日期
    if re.match(r"^\d{4}\s*年\s*\d{1,2}\s*月\s*\d{1,2}\s*日$", line):
        return "signature"

    # 第一条非空行 → 标题
    if is_first_content:
        return "title"

    # 一级标题：一、二、三、...
    if re.match(r"^[一二三四五六七八九十]、", line):
        return "heading1"

    # 二级标题：（一）（二）...
    if re.match(r"^[（(][一二三四五六七八九十\d]+[）)]", line):
        return "heading2"

    # 主送机关：以冒号结尾的短行
    # 排除含引导词的正文行（如「现将有关事项通知如下：」「具体要求如下：」）
    if line.endswith("：") and len(line) < 30:
        if not re.search(r"如下|以下|为$|是$", line):
            return "recipient"
        return "body"

    # 落款单位（短行，位于正文之后）
    # 无法精确判断，交给 body 处理
    return "body"


# ============================================================
# docx 写入器
# ============================================================

def save_gongwen_to_docx(
    content: str,
    output_path: Path = None,
    title: str = None,
) -> Path:
    """
    将生成的公文正文渲染为 .docx 文件。

    Args:
        content:     公文正文（纯文本，由 Agent 生成）。
        output_path: 输出路径。留空则自动生成到 data/output/ 目录。
        title:       公文标题（用于文件命名）。留空则从 content 中提取。

    Returns:
        生成的 .docx 文件路径。
    """
    # ---- 确定输出路径 ----
    if output_path is None:
        output_dir = _DATA_OUTPUT_DIR
        output_dir.mkdir(parents=True, exist_ok=True)
        # 用标题 + 时间戳做文件名
        safe_title = (title or "公文")[:30]
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', safe_title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_dir / f"{safe_title}_{timestamp}.docx"

    # ---- 创建文档 ----
    doc = Document()

    # ---- 页面设置 ----
    section = doc.sections[0]
    section.page_width = Cm(21.0)    # A4
    section.page_height = Cm(29.7)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    # ---- 设置默认样式 ----
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = SIZE_BODY
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
    # 行距：固定 28 磅
    style.paragraph_format.line_spacing = Pt(28)

    # ---- 逐行渲染 ----
    lines = content.strip().split('\n')
    is_first_content = True
    pending_signatures = []  # 延迟处理的落款行

    for line in lines:
        stripped = line.strip()
        role = _classify_line(stripped, is_first_content)

        if role == "blank":
            doc.add_paragraph()
            continue

        if role == "title":
            # 公文标题：二号小标宋体，居中
            p = _add_formatted_para(doc, stripped, FONT_TITLE, SIZE_TITLE, bold=False)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_first_content = False
            continue

        if role == "subtitle":
            # 发文字号：居中
            p = _add_formatted_para(doc, stripped, FONT_BODY, SIZE_BODY)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if role == "recipient":
            # 主送机关：顶格
            _add_formatted_para(doc, stripped, FONT_BODY, SIZE_BODY)
            continue

        if role == "heading1":
            # 一级标题：黑体
            _add_formatted_para(doc, stripped, FONT_HEITI, SIZE_BODY, bold=False)
            continue

        if role == "heading2":
            # 二级标题：楷体
            _add_formatted_para(doc, stripped, FONT_KAITI, SIZE_BODY, bold=False)
            continue

        if role == "signature":
            # 落款日期 → 需要和前一行（发文机关）组合右对齐
            pending_signatures.append(stripped)
            continue

        if role == "attachment":
            # 附件说明
            _add_formatted_para(doc, stripped, FONT_BODY, SIZE_BODY)
            continue

        # ---- body：正文 ----
        # 在 render 正文之前，先把积累的落款处理掉
        if pending_signatures:
            _flush_signatures(doc, pending_signatures)
            pending_signatures = []

        # 正文段落：首行缩进 2 字符
        p = _add_formatted_para(doc, stripped, FONT_BODY, SIZE_BODY)
        p.paragraph_format.first_line_indent = Cm(0.85 * 2)  # ~2个中文字符
        is_first_content = False

    # ---- 处理末尾的落款 ----
    if pending_signatures:
        _flush_signatures(doc, pending_signatures)

    # ---- 保存 ----
    doc.save(str(output_path))
    logger.info(f"公文已保存为 .docx：{output_path} （{os.path.getsize(output_path)/1024:.1f} KB）")

    return output_path


# ============================================================
# 辅助函数
# ============================================================

def _add_formatted_para(
    doc: Document,
    text: str,
    font_name: str,
    font_size,
    bold: bool = False,
) -> "Paragraph":
    """添加一个带字体格式的段落。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    # 行距
    p.paragraph_format.line_spacing = Pt(28)
    return p


def _flush_signatures(doc: Document, lines: list):
    """
    处理落款区域（发文机关 + 成文日期），右对齐。
    如果有 2 行以上，最后一行是日期，其余是发文机关。
    """
    # 加一个空行隔开正文
    doc.add_paragraph()

    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(line)
        run.font.name = FONT_BODY
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        run.font.size = SIZE_BODY
        p.paragraph_format.line_spacing = Pt(28)
