"""
生成 4 篇标准公文范文（.docx），用于测试知识库构建流水线。
公文类型覆盖：请示、通知、报告、函。
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUTPUT_DIR = Path(__file__).resolve().parent.parent / "data" / "raw_docs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def set_font(run, name_cn="仿宋", size=16, bold=False):
    """设置中文字体（python-docx 中文字体需同时设置西文和东亚字体）。"""
    run.font.name = name_cn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)
    run.font.size = Pt(size)
    run.bold = bold


def add_body(doc, text: str):
    """添加正文段落（首行缩进 2 字符）。"""
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, "仿宋", 16)
        if line.strip():
            p.paragraph_format.first_line_indent = Cm(0.85)


def add_title(doc, text: str):
    """添加公文标题（居中、二号小标宋体）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "方正小标宋简体", 22, bold=True)


def add_doc_number(doc, text: str):
    """添加发文字号（居中）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "仿宋", 16)


def add_recipient(doc, text: str):
    """添加主送机关。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, "仿宋", 16)


def add_signature(doc, org: str, date_str: str):
    """添加落款（右对齐）。"""
    doc.add_paragraph()
    for text in (org, date_str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(text)
        set_font(run, "仿宋", 16)


# ============================================================
# 范文 1：请示
# ============================================================
doc1 = Document()
doc1.styles['Normal'].font.name = '仿宋'
doc1.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
add_title(doc1, "关于申请购置办公服务器的请示")
add_doc_number(doc1, "X信发〔2025〕18号")
add_recipient(doc1, "公司领导：")
add_body(doc1, """为保障公司业务系统稳定运行，我部门现有一批服务器已连续运行超过五年，设备老化严重，频繁出现硬件故障。经统计，2025年第一季度共发生硬件故障12次，导致系统中断服务累计超过40小时，严重影响正常业务开展。

根据信息化建设三年规划（2024-2026年）的总体部署，2025年度计划采购3台高性能服务器用于核心业务系统升级。经前期市场调研和技术评估，拟采购配置如下：

一、采购设备清单
（一）应用服务器2台：品牌型号为联想 ThinkSystem SR650，配置CPU Xeon Gold 5418Y × 2、内存256GB、硬盘SSD 960GB × 4。
（二）数据库服务器1台：品牌型号为浪潮 NF5280M7，配置CPU Xeon Gold 6426Y × 2、内存512GB、硬盘SSD 1.92TB × 6。

二、预算测算
经询价三家以上供应商，设备购置总预算约为人民币肆拾捌万元整，具体明细如下：
1. 应用服务器2台：单价约15万元/台，合计30万元；
2. 数据库服务器1台：单价约18万元/台，合计18万元。
以上费用拟从2025年度信息化专项经费中列支。

三、实施计划
拟于2025年5月完成招标采购，6月完成设备到货安装，7月完成系统迁移和调试，8月正式投入运行。

妥否，请批示。""")
doc1.add_paragraph()
add_body(doc1, "附件：1. 服务器配置详细参数表\n      2. 供应商报价对比表\n      3. 信息化专项经费预算表")
add_signature(doc1, "信息技术部", "2025年4月15日")
doc1.save(OUTPUT_DIR / "请示-申请购置办公服务器.docx")
print("范文1：请示 ✓")


# ============================================================
# 范文 2：通知
# ============================================================
doc2 = Document()
doc2.styles['Normal'].font.name = '仿宋'
doc2.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
add_title(doc2, "关于召开2025年安全生产工作会议的通知")
add_doc_number(doc2, "X安委发〔2025〕6号")
add_recipient(doc2, "各部门、各分公司：")
add_body(doc2, """为深入贯彻落实上级关于安全生产工作的决策部署，总结2024年度安全生产工作情况，分析当前面临的形势和存在的问题，安排部署2025年安全生产重点任务，经公司研究决定，召开2025年安全生产工作会议。现将有关事项通知如下：

一、会议时间
2025年3月20日（星期四）上午9:00，会期半天。

二、会议地点
公司总部三楼多功能会议厅。

三、参会人员
（一）公司领导班子成员；
（二）各部门负责人、安全管理员；
（三）各分公司主要负责人及安全生产分管领导；
（四）2024年度安全生产先进单位和个人代表。

四、会议议程
（一）总结2024年度安全生产工作，部署2025年重点任务；
（二）通报2024年度安全生产考核结果；
（三）表彰2024年度安全生产先进单位和个人；
（四）签订2025年度安全生产目标责任书；
（五）公司主要领导讲话。

五、有关要求
（一）请各参会单位于3月17日前将参会人员名单报送至安全管理部。
（二）各分公司负责人因故不能参会的，须书面请假并经分管领导批准。
（三）请参会人员提前10分钟入场，遵守会场纪律。
（四）会议期间请将手机调至静音状态。

联系人：张明，电话：010-88886666，邮箱：zhangming@example.com。""")
add_signature(doc2, "XX公司安全生产委员会", "2025年3月10日")
doc2.save(OUTPUT_DIR / "通知-安全生产工作会议.docx")
print("范文2：通知 ✓")


# ============================================================
# 范文 3：报告
# ============================================================
doc3 = Document()
doc3.styles['Normal'].font.name = '仿宋'
doc3.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
add_title(doc3, "关于2024年度信息化建设工作情况的报告")
add_recipient(doc3, "市工业和信息化局：")
add_body(doc3, """2024年，我公司认真贯彻落实国家关于加快信息化建设的决策部署，按照《XX市数字化转型升级三年行动方案（2023-2025年）》的要求，扎实推进信息化建设各项工作，取得阶段性成效。现将有关情况报告如下：

一、主要工作进展
（一）网络基础设施建设。完成公司总部及5个分公司办公网络升级改造，实现万兆骨干网络全覆盖，网络可用率达到99.9%。
（二）核心业务系统建设。ERP系统完成二期开发并上线运行，实现了财务、采购、库存、销售全流程数字化管理。OA协同办公系统完成升级，新增移动审批、电子签章、智能公文三大模块。
（三）网络安全防护。完成等保2.0三级测评，部署下一代防火墙、入侵检测系统和终端安全管理系统，全年未发生重大网络安全事件。
（四）数据治理工作。搭建企业数据中台，完成核心业务数据资产梳理和标准化治理，编制发布《数据资产目录（V1.0）》。

二、存在的问题和不足
一是部分老旧系统整合难度较大，存在信息孤岛现象。二是信息化专业人才储备不足，高端复合型人才紧缺。三是数据价值挖掘利用不够充分，尚未形成有效的数据驱动决策机制。

三、2025年工作计划
（一）持续推进ERP三期建设，重点完善生产制造和质量管理模块。
（二）启动智能制造平台建设，推动生产线数字化改造。
（三）加强信息化人才队伍建设，计划引进和培养专业技术人才15名。
（四）深化数据应用，建设经营管理驾驶舱，支撑管理决策。

特此报告。""")
add_signature(doc3, "XX公司", "2025年1月15日")
doc3.save(OUTPUT_DIR / "报告-年度信息化建设工作情况.docx")
print("范文3：报告 ✓")


# ============================================================
# 范文 4：函
# ============================================================
doc4 = Document()
doc4.styles['Normal'].font.name = '仿宋'
doc4.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
add_title(doc4, "关于商请协助开展数据共享对接工作的函")
add_doc_number(doc4, "X数函〔2025〕12号")
add_recipient(doc4, "市交通运输局：")
add_body(doc4, """为加快推进我市智慧城市建设，实现跨部门数据互通共享，根据《XX市政务数据资源共享管理办法》和全市政务数据共享工作推进会的有关要求，我局拟与贵局开展交通领域数据共享对接工作。现就有关事项商洽如下：

一、数据共享需求
根据智慧城市一体化平台建设需要，拟共享以下数据资源：
（一）实时公交运行数据：包括公交车辆GPS定位、到站预测、线路站点信息等。
（二）交通流量监测数据：主要道路和交叉路口的车流量、平均车速等统计数据。
（三）交通事件信息：道路施工、交通管制、事故预警等实时事件数据。
（四）停车场信息：公共停车场位置、泊位总数、实时空余泊位等。

二、技术对接方案
拟采用市级政务数据共享交换平台作为数据传输通道，通过API接口方式实现数据实时同步。数据交换格式采用JSON，传输协议采用HTTPS，安全认证采用OAuth 2.0。我局已预留数据接收接口，接口规范详见附件。

三、数据安全管理
双方严格遵守《数据安全法》和《个人信息保护法》相关规定，建立数据使用审批制度，确保数据仅用于政务服务和城市管理领域，不得向第三方泄露或挪作他用。

四、工作安排
建议于2025年3月由双方技术团队召开对接协调会，明确具体实施方案和时间节点。争取在2025年6月底前完成首批数据对接。

以上事项，请予研究函复为盼。""")
doc4.add_paragraph()
add_body(doc4, "附件：1. 数据共享需求清单\n      2. API接口技术规范\n      3. 数据安全承诺书模板")
add_signature(doc4, "XX市数据资源管理局", "2025年3月5日")
doc4.save(OUTPUT_DIR / "函-商请协助数据共享对接.docx")
print("范文4：函 ✓")

print(f"\n全部范文已保存到 {OUTPUT_DIR}")
for f in sorted(OUTPUT_DIR.glob("*.docx")):
    print(f"  {f.name}  ({f.stat().st_size / 1024:.1f} KB)")
